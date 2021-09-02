from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_doc(name, value, service):
    document = Document()
    today_str = datetime.now().strftime("%d/%m/%Y")

    today_str_hour = datetime.now().strftime("%d%m%Y-%H%M%S")


    document.add_heading("Recibo")

    document.add_paragraph("Recebemos de ")

    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(name)
    r1.italic = True
    
    document.add_paragraph("A quantia de ")

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"R$ {value}")
    r2.italic = True
    
    document.add_paragraph("Referente ao serviço de ")

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run(service)
    r3.italic = True

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(today_str)
    r4.bold = True

    document.save(f"Recibo {today_str_hour}.docx")
